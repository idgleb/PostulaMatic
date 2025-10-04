#!/usr/bin/env python
"""
Parser especializado para archivos DOCX.
Extrae texto de párrafos, tablas, headers, footers y elementos gráficos.
"""

import logging
import os
import re
from typing import Dict

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parser especializado para archivos DOCX."""

    def __init__(self):
        self.supported_extensions = [".docx"]

    def is_supported(self, file_path: str) -> bool:
        """Verifica si el archivo es soportado por este parser."""
        if not Document:
            return False

        extension = os.path.splitext(file_path)[1].lower()
        return extension in self.supported_extensions

    def parse_cv(self, file_path: str) -> Dict:
        """
        Parsea un archivo DOCX y extrae todo el texto disponible.

        Args:
            file_path: Ruta al archivo DOCX

        Returns:
            Dict con 'text' y 'warning_message' si aplica
        """
        if not self.is_supported(file_path):
            return {"text": "", "warning_message": "Formato DOCX no soportado"}

        try:
            doc = Document(file_path)

            # Extraer texto de diferentes fuentes
            text_parts = []

            # 1. Párrafos principales
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # 2. Tablas
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)

            # 3. Headers y footers
            for section in doc.sections:
                # Header
                if section.header:
                    header_text = self._extract_header_footer_text(section.header)
                    if header_text:
                        text_parts.append(header_text)

                # Footer
                if section.footer:
                    footer_text = self._extract_header_footer_text(section.footer)
                    if footer_text:
                        text_parts.append(footer_text)

            # Unir todo el texto
            full_text = "\n\n".join(text_parts)

            # Limpiar y normalizar
            clean_text = self._clean_docx_text(full_text)

            # Verificar si hay suficiente contenido
            if len(clean_text.strip()) < 50:
                return {
                    "text": clean_text,
                    "warning_message": "Texto extraído muy corto. Verifica que el documento contenga texto.",
                }

            return {"text": clean_text, "warning_message": None}

        except Exception as e:
            logger.error(f"Error parseando DOCX {file_path}: {e}")
            return {"text": "", "warning_message": f"Error parseando DOCX: {str(e)}"}

    def _extract_table_text(self, table: Table) -> str:
        """Extrae texto de una tabla DOCX."""
        table_rows = []

        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)

            if row_cells:
                table_rows.append("\t".join(row_cells))

        return "\n".join(table_rows) if table_rows else ""

    def _extract_header_footer_text(self, header_footer) -> str:
        """Extrae texto de headers y footers."""
        text_parts = []

        for paragraph in header_footer.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        return "\n".join(text_parts) if text_parts else ""

    def _clean_docx_text(self, text: str) -> str:
        """
        Limpia el texto extraído del DOCX.

        Args:
            text: Texto crudo del DOCX

        Returns:
            Texto limpio y normalizado
        """
        if not text:
            return ""

        # Limpiar caracteres de control
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

        # Normalizar espacios múltiples
        text = re.sub(r"\s+", " ", text)

        # Limpiar líneas vacías múltiples
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)

        return text.strip()

    def get_supported_formats(self) -> list:
        """Retorna los formatos soportados por este parser."""
        return self.supported_extensions
