"""
Servicio para parsear CVs en formato PDF y DOCX.
Extrae el texto limpio y lo normaliza para procesamiento posterior.
"""
import os
import logging
from typing import Optional, Dict, Any
from pathlib import Path

# PDF parsing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class CVParserError(Exception):
    """Excepción personalizada para errores del parser de CV."""
    pass


class CVParser:
    """Parser principal para CVs en diferentes formatos."""
    
    def __init__(self):
        self.supported_formats = []
        if PDF_AVAILABLE:
            self.supported_formats.append('.pdf')
        if DOCX_AVAILABLE:
            self.supported_formats.append('.docx')
    
    def parse_cv(self, file_path: str) -> Dict[str, Any]:
        """
        Parsea un archivo de CV y extrae el texto.
        
        Args:
            file_path: Ruta al archivo de CV
            
        Returns:
            Dict con 'text', 'format', 'word_count', 'pages' (si aplica)
            
        Raises:
            CVParserError: Si el archivo no se puede procesar
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise CVParserError(f"El archivo {file_path} no existe")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise CVParserError(
                f"Formato {file_extension} no soportado. "
                f"Formatos soportados: {', '.join(self.supported_formats)}"
            )
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension == '.docx':
                return self._parse_docx(file_path)
        except Exception as e:
            logger.error(f"Error parseando CV {file_path}: {e}")
            raise CVParserError(f"Error procesando el archivo: {str(e)}")
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parsea un archivo PDF con extracción mejorada para estructuras complejas."""
        if not PDF_AVAILABLE:
            raise CVParserError("PyPDF2 no está disponible para procesar PDFs")
        
        text_content = []
        pages = 0
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            pages = len(pdf_reader.pages)
            
            for page_num in range(pages):
                try:
                    page = pdf_reader.pages[page_num]
                    
                    # 1. Extracción básica de texto
                    basic_text = page.extract_text()
                    if basic_text.strip():
                        text_content.append(basic_text)
                    
                    # 2. Extracción mejorada para estructuras complejas
                    try:
                        enhanced_text = self._extract_enhanced_pdf_text(page)
                        if enhanced_text and enhanced_text != basic_text:
                            # Si la extracción mejorada es diferente, usarla
                            text_content[-1] = enhanced_text
                    except Exception as e:
                        logger.warning(f"Extracción mejorada falló en página {page_num + 1}: {e}")
                        continue
                        
                except Exception as e:
                    logger.warning(f"Error extrayendo página {page_num + 1}: {e}")
                    continue
        
        full_text = '\n\n'.join(text_content)
        
        # Normalizar el texto extraído
        normalized_text = self._normalize_text(full_text)
        
        # Detectar si el texto extraído es muy poco (posible problema de parsing)
        word_count = len(normalized_text.split())
        char_count = len(normalized_text.strip())
        
        # Si hay muy poco texto extraído, agregar advertencia
        warning_message = ""
        if pages > 0 and word_count < 10:
            warning_message = f"⚠️ ADVERTENCIA: Se extrajo muy poco texto ({word_count} palabras) de {pages} página(s). El archivo puede estar corrupto, ser una imagen escaneada, o tener protección."
        elif char_count < 50:
            warning_message = f"⚠️ ADVERTENCIA: Texto extraído muy corto ({char_count} caracteres). Verifica que el archivo sea un CV válido."
        
        return {
            'text': normalized_text,
            'format': 'pdf',
            'word_count': word_count,
            'pages': pages,
            'warning': warning_message,
            'extraction_method': 'enhanced_pdf'
        }
    
    def _extract_enhanced_pdf_text(self, page) -> str:
        """Extrae texto de manera mejorada de una página PDF."""
        try:
            # Intentar extraer texto con diferentes métodos
            text_parts = []
            
            # Método 1: Extracción básica
            basic_text = page.extract_text()
            if basic_text.strip():
                text_parts.append(basic_text)
            
            # Método 2: Extraer texto por bloques para preservar estructura
            try:
                # Intentar acceder a la estructura interna del PDF
                if hasattr(page, 'get_contents'):
                    contents = page.get_contents()
                    if contents:
                        # Procesar contenido del PDF
                        content_text = self._process_pdf_contents(contents)
                        if content_text:
                            text_parts.append(content_text)
            except Exception:
                pass
            
            # Método 3: Extraer texto con coordenadas para estructuras tabulares
            try:
                if hasattr(page, 'extract_text_simple'):
                    # Usar método alternativo si está disponible
                    alt_text = page.extract_text_simple()
                    if alt_text and alt_text != basic_text:
                        text_parts.append(alt_text)
            except Exception:
                pass
            
            # Combinar todos los métodos y limpiar duplicados
            combined_text = '\n'.join(text_parts)
            
            # Limpiar y reorganizar texto fragmentado
            cleaned_text = self._clean_fragmented_text(combined_text)
            
            return cleaned_text
            
        except Exception as e:
            logger.warning(f"Error en extracción mejorada: {e}")
            return basic_text
    
    def _process_pdf_contents(self, contents) -> str:
        """Procesa el contenido interno del PDF para extraer texto."""
        try:
            # Convertir contenido a texto si es posible
            if hasattr(contents, 'get_data'):
                data = contents.get_data()
                if isinstance(data, bytes):
                    # Intentar decodificar como texto
                    try:
                        return data.decode('utf-8', errors='ignore')
                    except:
                        return data.decode('latin-1', errors='ignore')
            return ""
        except Exception:
            return ""
    
    def _clean_fragmented_text(self, text: str) -> str:
        """Limpia y reorganiza texto fragmentado de PDFs complejos, filtrando metadatos PDF."""
        if not text:
            return ""
        
        # Filtrar metadatos PDF comunes
        pdf_metadata_patterns = [
            r'0\s+1\s+-1\s+0\s+\d+\s+0\s+cm',  # Transformaciones PDF
            r'BT.*?ET',  # Bloques de texto PDF
            r'q\s+.*?\s+Q',  # Estados de gráficos
            r'/\w+\s+\d+\s+Tf',  # Fuentes PDF
            r'\d+\s+\d+\s+Td',  # Posicionamiento de texto
            r'\d+\s+\d+\s+Tm',  # Matrices de transformación
            r'\[\]\s+\d+\s+d',  # Patrones de línea
            r'\d+\s+w\s+\d+\s+J',  # Grosor y estilo de línea
            r'\d+\s+\d+\s+\d+\s+rg',  # Colores RGB
            r'\d+\s+\d+\s+m\s+\d+\s+\d+\s+l',  # Líneas
            r'\d+\s+\d+\s+\d+\s+-\d+\s+re',  # Rectángulos
            r'f\s+.*?f',  # Rellenos
            r'S\s+.*?S',  # Trazos
        ]
        
        import re
        
        # Aplicar filtros para remover metadatos PDF
        cleaned_text = text
        for pattern in pdf_metadata_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.DOTALL)
        
        # Limpiar líneas vacías y espacios múltiples
        lines = cleaned_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 2:  # Solo líneas con contenido sustancial
                # Limpiar caracteres de control
                line = ''.join(char for char in line if ord(char) >= 32 or char in '\n\t')
                cleaned_lines.append(line)
        
        # Unir líneas y limpiar espacios múltiples
        cleaned_text = '\n'.join(cleaned_lines)
        cleaned_text = ' '.join(cleaned_text.split())  # Normalizar espacios
        
        return cleaned_text
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parsea un archivo DOCX con extracción completa de estructuras complejas."""
        if not DOCX_AVAILABLE:
            raise CVParserError("python-docx no está disponible para procesar DOCX")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # 1. Extraer texto de párrafos normales
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 2. Extraer texto de tablas (muy importante para documentos financieros)
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_content.append(table_text)
            
            # 3. Extraer texto de encabezados y pies de página
            for section in doc.sections:
                # Encabezados
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
                
                # Pies de página
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
            
            # 4. Extraer texto de formas y cuadros de texto (si están disponibles)
            try:
                # Intentar extraer texto de elementos gráficos
                for element in doc.element.body:
                    if hasattr(element, 'text'):
                        if element.text and element.text.strip():
                            text_content.append(element.text)
            except:
                # Si falla, continuar sin error
                pass
            
            # 5. Unir todo el texto extraído
            full_text = '\n'.join(text_content)
            
            # Normalizar el texto extraído
            normalized_text = self._normalize_text(full_text)
            
            # Detectar si el texto extraído es muy poco
            word_count = len(normalized_text.split())
            char_count = len(normalized_text.strip())
            
            # Si hay muy poco texto extraído, agregar advertencia
            warning_message = ""
            if word_count < 10:
                warning_message = f"⚠️ ADVERTENCIA: Se extrajo muy poco texto ({word_count} palabras). El archivo puede estar corrupto, vacío, o tener protección."
            elif char_count < 50:
                warning_message = f"⚠️ ADVERTENCIA: Texto extraído muy corto ({char_count} caracteres). Verifica que el archivo sea un CV válido."
            
            return {
                'text': normalized_text,
                'format': 'docx',
                'word_count': word_count,
                'pages': 1,  # DOCX no tiene páginas definidas
                'warning': warning_message,
                'extraction_method': 'enhanced'  # Indicar que se usó extracción mejorada
            }
            
        except Exception as e:
            raise CVParserError(f"Error procesando archivo DOCX: {str(e)}")
    
    def _extract_table_text(self, table) -> str:
        """Extrae texto de una tabla DOCX de manera estructurada."""
        table_rows = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # Extraer texto de cada celda
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            
            if row_cells:
                # Unir celdas con tabulador para mantener estructura
                table_rows.append('\t'.join(row_cells))
        
        return '\n'.join(table_rows)
    
    def _normalize_text(self, text: str) -> str:
        """
        Normaliza el texto extraído para mejor procesamiento.
        Reconstruye texto fragmentado línea por línea de manera inteligente.
        Basado en el CV real del usuario.
        
        Args:
            text: Texto crudo extraído del CV
            
        Returns:
            Texto normalizado
        """
        if not text:
            return ""
        
        import re
        
        # Remover caracteres no imprimibles
        text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
        
        # Normalizar espacios múltiples
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Dividir en líneas
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            return ""
        
        reconstructed_lines = []
        current_section = ""
        current_paragraph = ""
        
        for i, line in enumerate(lines):
            # Detectar secciones principales
            is_section_header = line in [
                'PERFIL PROFESIONAL', 'EXPERIENCIA LABORAL', 'EDUCACIÓN', 
                'PROYECTOS DESTACADOS', 'HABILIDADES', 'IDIOMAS', 'CERTIFICACIONES'
            ]
            
            # Detectar títulos de trabajo/empresa
            is_job_title = re.match(r'^\d{4}', line) or 'Desarrollador' in line or 'Developer' in line
            
            # Detectar URLs y emails
            is_url_or_email = '@' in line or 'http' in line or '.com' in line or 'github.com' in line or 'linkedin.com' in line
            
            # Detectar si es el final de una oración
            ends_with_period = line.endswith('.') or line.endswith('!') or line.endswith('?')
            
            # Si es una sección principal, mantenerla separada
            if is_section_header:
                if current_paragraph:
                    reconstructed_lines.append(current_paragraph.strip())
                    current_paragraph = ""
                reconstructed_lines.append(line)
                current_section = line
            
            # Si es un título de trabajo, mantenerlo separado
            elif is_job_title:
                if current_paragraph:
                    reconstructed_lines.append(current_paragraph.strip())
                    current_paragraph = ""
                reconstructed_lines.append(line)
            
            # Si es URL o email, mantenerlo separado
            elif is_url_or_email:
                if current_paragraph:
                    reconstructed_lines.append(current_paragraph.strip())
                    current_paragraph = ""
                reconstructed_lines.append(line)
            
            # Si termina con punto, completar párrafo
            elif ends_with_period:
                current_paragraph += " " + line if current_paragraph else line
                reconstructed_lines.append(current_paragraph.strip())
                current_paragraph = ""
            
            # Si es una línea corta (probablemente fragmentada), continuar construyendo
            elif len(line.split()) <= 3 and not line.isupper():
                current_paragraph += " " + line if current_paragraph else line
            
            # Si es una línea normal, continuar construyendo
            else:
                current_paragraph += " " + line if current_paragraph else line
        
        # Agregar el último párrafo si queda algo
        if current_paragraph:
            reconstructed_lines.append(current_paragraph.strip())
        
        # Unir líneas y limpiar
        text = '\n'.join(reconstructed_lines)
        
        # Limpiar saltos de línea múltiples
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        
        # Remover líneas vacías al inicio y final
        text = text.strip()
        
        return text
    
    def get_supported_formats(self) -> list:
        """Retorna la lista de formatos soportados."""
        return self.supported_formats.copy()
    
    def is_supported(self, file_path: str) -> bool:
        """Verifica si el archivo es de un formato soportado."""
        return Path(file_path).suffix.lower() in self.supported_formats


# Instancia global del parser
cv_parser = CVParser()
